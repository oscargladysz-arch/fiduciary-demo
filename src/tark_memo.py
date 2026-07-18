"""
Tark decision memo generator (M5)
=================================
Renders a per-product Word decision memo from the evaluated data: regulatory
basis, six-factor findings, benchmark selection with the FULL rejection log
(or the escalation, for the fail case), and a provenance appendix. The memo is
the artifact a fiduciary files; the rejection log is half its legal value.

Run:  python src/tark_memo.py            -> data/memos/<key>_decision_memo.docx
Anonymization: only the plan's display label ever appears (test-enforced).
"""
from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from tark_data import (DATA, FACTORS, cells_by_factor, load_anchor_plan,
                       load_products, status_kind)

RULE = ("DOL proposed rule, Fiduciary Duties in Selecting Designated "
        "Investment Alternatives, 91 FR 16088 (Mar. 31, 2026), RIN 1210-AC38")

RULE_PARAS = [
    ("This memo documents an evaluation of the product below as a potential "
     "designated investment alternative (DIA) for the plan, structured on the "
     "six factors of the proposed rule: performance, fees and expenses, "
     "liquidity, valuation, performance benchmarks, and complexity. The "
     "proposed safe harbor attaches to a documented, objective, thorough and "
     "analytical process; this memo and its underlying evidence files "
     "constitute that record."),
    ("On benchmarks, the proposal requires comparison against a meaningful "
     "benchmark and acknowledges that no single benchmark is meaningful for "
     "every DIA; where none exists, the history of a similar type of "
     "investment may serve. The benchmark section below therefore documents "
     "the selection AND the rejections: every candidate considered, its "
     "score, and the true reason it was or was not chosen. Note: the pleading "
     "standard for benchmark-based claims is before the Supreme Court in "
     "Anderson v. Intel (argument expected October Term 2026); this memo's "
     "approach is designed to be defensible under either outcome."),
]


def _cell_lines(product: dict, factor_label: str, limit: int = 4) -> str:
    picks = []
    for cid, cell in cells_by_factor(product)[factor_label]:
        if status_kind(cell.get("status", "")) in ("extracted", "verified",
                                                   "computed", "partial"):
            v = (cell.get("value") or "").strip()
            if v:
                picks.append(f"{cid}: {v[:220]}")
        if len(picks) >= limit:
            break
    return "\n".join(picks) if picks else "No cells evaluated yet — pending extraction."


def _set_letter(doc: Document) -> None:
    s = doc.sections[0]
    s.page_width, s.page_height = Inches(8.5), Inches(11)
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(s, m, Inches(1))


def build_memo(key: str) -> Path:
    products = load_products()
    p = products[key]
    anchor = load_anchor_plan()
    sel_path = DATA / "benchmarks" / f"{key}_selection.json"
    sel = json.loads(sel_path.read_text()) if sel_path.exists() else None

    doc = Document()
    _set_letter(doc)
    doc.styles["Normal"].font.size = Pt(10.5)

    doc.add_heading("Designated Investment Alternative Evaluation — "
                    "Decision Memo", level=0)
    sub = doc.add_paragraph()
    sub.add_run(f"Plan: {anchor['display_label']}\n").bold = True
    sub.add_run(f"Product: {p['fund_name']} ({p['wrapper']}, CIK {p['cik']})\n")
    sub.add_run(f"Date: {date.today().isoformat()}    Status: DRAFT — demo "
                "build; cells marked extracted-unverified pend independent "
                "verification.")

    doc.add_heading("Regulatory basis", level=1)
    doc.add_paragraph(RULE)
    for para in RULE_PARAS:
        doc.add_paragraph(para)

    doc.add_heading("Six-factor findings (summary)", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Factor", "Key findings (cell: value)"
    for n, label in FACTORS.items():
        row = table.add_row().cells
        row[0].text = f"{n} — {label}"
        row[1].text = _cell_lines(p, label)
    for row in table.rows:
        row.cells[0].width, row.cells[1].width = Inches(1.4), Inches(5.1)

    doc.add_heading("Benchmark selection and justification", level=1)
    if sel is None:
        doc.add_paragraph("Engine profile pending for this product — "
                          "extraction depth required before selection.")
    else:
        if sel.get("escalation"):
            doc.add_heading("ESCALATION — no meaningful benchmark "
                            "constructible", level=2)
            doc.add_paragraph(sel["escalation"])
            doc.add_paragraph(
                "Under the proposal's own terms, a benchmark that is not "
                "meaningful cannot support the comparison; proceeding without "
                "one documented here would undercut the safe harbor. "
                "Recommended action: do not proceed pending the data steps "
                "above; retain this memo as the record of the determination.")
        for slot, badge in (("primary", "Primary"), ("secondary", "Secondary")):
            s = sel.get(slot)
            if not s:
                continue
            doc.add_heading(f"{badge}: {s['candidate']} — score "
                            f"{s['score']}/{s['max']}", level=2)
            comp = s.get("comparison")
            if comp:
                doc.add_paragraph(
                    f"Window {comp['window']}: fund {comp['fund_ann_pct']}%/yr "
                    f"vs benchmark {comp['index_ann_pct']}%/yr; "
                    f"KS-PME {comp['ks_pme']}; Direct Alpha "
                    f"{comp['direct_alpha_pct']}%/yr. Disclosure: PME and "
                    "alpha computed on appraisal-lagged NAVs are "
                    "window-sensitive and can be smoothing-flattered; "
                    "conclusions should be read with the methodology's "
                    "window-sensitivity analysis.")
            for r in s["reasons"]:
                doc.add_paragraph(r, style="List Bullet")

        doc.add_heading("Rejection log (candidates considered and not "
                        "selected)", level=2)
        rt = doc.add_table(rows=1, cols=3)
        rt.style = "Table Grid"
        h = rt.rows[0].cells
        h[0].text, h[1].text, h[2].text = "Candidate", "Score", "Reason"
        for r in sel["rejected"]:
            c = rt.add_row().cells
            c[0].text = r["candidate"]
            c[1].text = f"{r['score']}/{r['max']}"
            c[2].text = r["rejection"]
        for row in rt.rows:
            row.cells[0].width = Inches(2.4)
            row.cells[1].width = Inches(0.8)
            row.cells[2].width = Inches(3.3)

    doc.add_heading("Provenance", level=1)
    counts: dict[str, int] = {}
    for cell in p["cells"].values():
        counts[status_kind(cell.get("status", ""))] = \
            counts.get(status_kind(cell.get("status", "")), 0) + 1
    doc.add_paragraph(
        "Every populated cell carries its source document, section, quote, "
        "extractor and verifier in data/evidence/. Current cell status for "
        "this product: "
        + ", ".join(f"{k}: {v}" for k, v in sorted(counts.items())) + ". "
        "Cells marked extracted-unverified or computed await independent "
        "verification; verified cells have been independently re-checked.")
    doc.add_paragraph(f"Data sources include SEC EDGAR filings (see "
                      f"data/manifest.csv) and DOL EBSA Form 5500 bulk data. "
                      f"{anchor['anonymization_rule']}")

    sig = doc.add_paragraph()
    sig.add_run("\nPrepared by: ______________________    "
                "Reviewed by: ______________________    "
                "Date: ____________")
    sig.alignment = WD_ALIGN_PARAGRAPH.LEFT

    out = DATA / "memos"
    out.mkdir(exist_ok=True)
    path = out / f"{key}_decision_memo.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    for key in load_products():
        print("wrote", build_memo(key))
