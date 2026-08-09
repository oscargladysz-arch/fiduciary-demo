"""Memo regression checks (M5). Run: python src/test_memo.py"""
import sys
from pathlib import Path
from docx import Document

BASE = Path(__file__).resolve().parents[1]
FAILS = []
def check(name, cond):
    print(f"[{'PASS' if cond else 'FAIL'}] {name}")
    if not cond: FAILS.append(name)

def text_of(p):
    d = Document(p)
    parts = [para.text for para in d.paragraphs]
    for t in d.tables:
        for r in t.rows:
            for c in r.cells: parts.append(c.text)
    return "\n".join(parts).lower()

keys = ["breit","cliffwater_cclfx","dxyz","hl_paf","kkr_kpec","stepstone_spm"]
texts = {}
for k in keys:
    p = BASE / "data" / "memos" / f"{k}_decision_memo.docx"
    check(f"{k}: memo exists", p.exists())
    if p.exists():
        texts[k] = text_of(p)
        check(f"{k}: anonymization holds", "spotify" not in texts[k])
        check(f"{k}: rule cited", "91 fr 16088" in texts[k])
check("cclfx: PME + CDLI rejection in memo",
      "ks-pme 1.2532" in texts["cliffwater_cclfx"] and "cliffwater direct lending index" in texts["cliffwater_cclfx"])
check("dxyz: escalation variant", "escalation" in texts["dxyz"] and "no meaningful benchmark" in texts["dxyz"])
check("kpec: K-1 finding in memo", "schedule k-1" in texts["kkr_kpec"])
check("paf: window-sensitivity disclosure", "window-sensitive" in texts["hl_paf"])
print(f"\n{len(FAILS)} failure(s)." if FAILS else "\nAll memo checks pass.")
sys.exit(1 if FAILS else 0)
